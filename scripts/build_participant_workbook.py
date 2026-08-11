#!/usr/bin/env python3
"""Build the participant workbook from the canonical Markdown source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = RGBColor(0x1D, 0x26, 0x25)
DARK_GREEN = RGBColor(0x12, 0x3C, 0x3A)
ORANGE = RGBColor(0xE7, 0x66, 0x3E)
GOLD = RGBColor(0xD4, 0xA7, 0x2C)
MUTED = RGBColor(0x5D, 0x6A, 0x68)
PALE_GREEN = "E5EAE6"
PALE_ORANGE = "F8E5DE"
LIGHT_RULE = "D8DED9"


def set_font(run, name: str, size: float, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, edge: str, color: str, size=6, space=3):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_font(run, "Calibri", 9, MUTED)


def add_numbering_definition(document: Document, marker: str, num_format: str, start=1):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:start")
    start_node.set(qn("w:val"), str(start))
    level.append(start_node)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), num_format)
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), marker)
    level.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "E7663E")
    r_pr.append(color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def configure_page(section, top_margin=1.0):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(top_margin)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def populate_footer(footer):
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(fp, before=3, after=0, line=1)
    add_page_field(fp)


def configure_body_header_footer(section):
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def configure_document(document: Document):
    document.settings.odd_and_even_pages_header_footer = True
    cover_section = document.sections[0]
    configure_page(cover_section)
    cover_section.different_first_page_header_footer = False

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, DARK_GREEN, 18, 10),
        ("Heading 2", 13, DARK_GREEN, 14, 7),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x4A), 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Workbook Callout" not in styles:
        callout = styles.add_style("Workbook Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Workbook Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = DARK_GREEN
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2

def add_cover(document: Document):
    for _ in range(4):
        p = document.add_paragraph()
        set_spacing(p, after=12)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(kicker, after=16)
    run = kicker.add_run("ПРАКТИЧЕСКОЕ РУКОВОДСТВО")
    set_font(run, "Calibri", 10, ORANGE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, after=8, line=1)
    run = title.add_run("Нейросети для специалиста")
    set_font(run, "Georgia", 30, DARK_GREEN, bold=True)

    year = document.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(year, after=28, line=1)
    run = year.add_run("2026")
    set_font(run, "Georgia", 22, ORANGE, bold=True)

    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(lead, after=32, line=1.25)
    run = lead.add_run(
        "Одна повторяемая задача. Три реальных повтора. "
        "Одно честное решение по результату."
    )
    set_font(run, "Calibri", 13.5, INK)

    for label, value in [
        ("ШАГ 01", "Выберите задачу"),
        ("ШАГ 02", "Соберите бриф и проверку"),
        ("ШАГ 03", "Сравните полный цикл"),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, after=8, line=1.15)
        r1 = p.add_run(f"{label}  ")
        set_font(r1, "Calibri", 9.5, GOLD, bold=True)
        r2 = p.add_run(value)
        set_font(r2, "Calibri", 11, DARK_GREEN, bold=True)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(48)
    set_spacing(note, after=0, line=1.15)
    run = note.add_run(
        "Нейросеть готовит черновик. Факты, решения и ответственность "
        "остаются у человека."
    )
    set_font(run, "Calibri", 10, MUTED, italic=True)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section, top_margin=0.8)
    configure_body_header_footer(body_section)


def add_response_lines(document: Document, count=3):
    for _ in range(count):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(7)
        set_paragraph_border(p, "bottom", "B8C3BD", size=4, space=1)
        p.add_run(" ")


def add_callout(document: Document, text: str, fill=PALE_GREEN):
    p = document.add_paragraph(style="Workbook Callout")
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, "left", "E7663E", size=18, space=5)
    run = p.add_run(text)
    set_font(run, "Calibri", 10.5, DARK_GREEN)
    return p


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text.strip()


def build_from_markdown(document: Document, markdown: str):
    lines = markdown.splitlines()
    bullet_num_id = add_numbering_definition(document, "•", "bullet")
    check_num_id = add_numbering_definition(document, "☐", "bullet")
    number_num_id = add_numbering_definition(document, "%1.", "decimal")

    response_headings = {
        "Моя роль": 2,
        "Задача": 3,
        "Частота и время": 3,
        "Материал на входе": 3,
        "Ожидаемый результат": 3,
        "Название сценария": 2,
        "Когда запускается": 2,
        "Вход": 3,
        "Инструкция": 3,
        "Выход": 2,
        "Проверка": 3,
        "Измерения": 3,
        "Решение": 2,
    }
    major_page_breaks = {
        "2. Бриф для ИИ: шесть опор",
        "3. Три готовых запроса",
        "4. Чек-лист проверки",
        "5. Три повтора",
        "6. Карточка рабочего процесса",
    }
    continuation_page_breaks = {
        "Ожидаемый результат",
        "Проверка",
    }

    i = 0
    skipped_title = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        i += 1
        if not stripped:
            continue

        if stripped.startswith("# ") or stripped.startswith("## Рабочая тетрадь"):
            skipped_title += 1
            continue

        if stripped.startswith("## "):
            title = clean_inline(stripped[3:])
            p = document.add_paragraph(title, style="Heading 1")
            if title in major_page_breaks:
                p.paragraph_format.page_break_before = True
            p.paragraph_format.keep_with_next = True
            set_paragraph_border(p, "bottom", "E7663E", size=6, space=4)
            continue

        if stripped.startswith("### "):
            title = clean_inline(stripped[4:])
            p = document.add_paragraph(title, style="Heading 2")
            if title in continuation_page_breaks:
                p.paragraph_format.page_break_before = True
            if title in response_headings:
                add_response_lines(document, response_headings[title])
            continue

        if stripped.startswith(">"):
            quote_lines = [clean_inline(stripped[1:].strip())]
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line.startswith(">"):
                    break
                quote_lines.append(clean_inline(next_line[1:].strip()))
                i += 1
            text = "\n".join(line for line in quote_lines if line)
            add_callout(document, text, PALE_ORANGE)
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            p = document.add_paragraph()
            apply_numbering(p, number_num_id)
            set_spacing(p, after=4, line=1.25)
            run = p.add_run(clean_inline(numbered.group(2)))
            set_font(run, "Calibri", 11, INK)
            continue

        if stripped.startswith("- "):
            text = clean_inline(stripped[2:])
            p = document.add_paragraph()
            apply_numbering(
                p,
                check_num_id
                if any(
                    word in text.lower()
                    for word in [
                        "сверен",
                        "существ",
                        "отвечает",
                        "подходит",
                        "обещан",
                        "разрешено",
                        "удалены",
                        "публикуются",
                        "повторяется",
                        "цифров",
                        "описать",
                        "проверить",
                        "ущерб",
                    ]
                )
                else bullet_num_id,
            )
            set_spacing(p, after=4, line=1.25)
            run = p.add_run(text)
            set_font(run, "Calibri", 11, INK)
            continue

        paragraph_lines = [stripped]
        while i < len(lines):
            next_line = lines[i].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("- ")
                or next_line.startswith(">")
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            i += 1
        text = clean_inline(" ".join(paragraph_lines))
        if text.startswith("Материал не является"):
            add_callout(document, text, PALE_ORANGE)
        else:
            p = document.add_paragraph()
            set_spacing(p, after=7, line=1.25)
            run = p.add_run(text)
            set_font(run, "Calibri", 11, INK)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    source = Path(args.source)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    document.core_properties.title = "Нейросети для специалиста - 2026"
    document.core_properties.subject = "Рабочая тетрадь участника"
    document.core_properties.author = "Alex Fintore"
    document.core_properties.keywords = "нейросети, работа, практикум, рабочая тетрадь"

    add_cover(document)
    markdown = source.read_text(encoding="utf-8")
    build_from_markdown(document, markdown)

    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
