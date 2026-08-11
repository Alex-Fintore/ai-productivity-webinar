#!/usr/bin/env python3
"""Build the post-webinar starter guide from a small JSON content source."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(0x08, 0x14, 0x2D)
INDIGO = RGBColor(0x4B, 0x57, 0xE8)
AQUA = RGBColor(0x15, 0x9E, 0x9A)
YELLOW = RGBColor(0xD9, 0x9A, 0x00)
CORAL = RGBColor(0xC4, 0x3A, 0x2F)
MUTED = RGBColor(0x4C, 0x58, 0x75)
LAVENDER = "F2F0FF"
CREAM = "FFFDF7"
WHITE = "FFFFFF"
RULE = "DADDF0"


def set_font(run, size=11, color=NAVY, bold=False, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url, color=INDIGO, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), str(color))
    r_pr.append(color_node)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")
    r_pr.append(size)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    link.append(run)
    paragraph._p.append(link)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_numbering(document, marker, num_format, color="159E9A"):
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", num_format), ("w:lvlText", marker), ("w:lvlJc", "left")):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    marker_color = OxmlElement("w:color")
    marker_color.set(qn("w:val"), color)
    r_pr.append(marker_color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def add_callout(document, text, fill=LAVENDER, color=NAVY):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_spacing(p, before=2, after=2, line=1.2)
    run = p.add_run(text)
    set_font(run, size=11.5, color=color, bold=True)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(document, text, level=1):
    p = document.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(document, text, bold_lead=None):
    p = document.add_paragraph()
    set_spacing(p)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, INDIGO, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [4680, 4680])
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell, text in zip(table.rows[0].cells, ("НЕЙРОСЕТИ ПРОСТЫМ ЯЗЫКОМ", "МАТЕРИАЛЫ ВЕБИНАРА")):
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1)
        set_font(p.add_run(text), size=8.5, color=MUTED, bold=True, name="Courier New")
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_spacing(p, after=0, line=1)
    add_page_field(p)


def add_cover(document, data):
    p = document.add_paragraph()
    set_spacing(p, before=18, after=4, line=1)
    set_font(p.add_run("МАТЕРИАЛЫ УЧАСТНИКА"), size=10, color=INDIGO, bold=True, name="Courier New")
    p = document.add_paragraph()
    set_spacing(p, after=6, line=0.95)
    set_font(p.add_run(data["title"]), size=31, color=NAVY, bold=True)
    p = document.add_paragraph()
    set_spacing(p, after=20, line=1.1)
    set_font(p.add_run(data["subtitle"]), size=15, color=INDIGO, bold=True)

    meta = document.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680])
    values = (("Формат", "памятка после вебинара"), ("Проверено", data["updated"]))
    for row, pair in zip(meta.rows, values):
        for cell, text in zip(row.cells, pair):
            shade_cell(cell, CREAM)
            p = cell.paragraphs[0]
            set_spacing(p, after=0, line=1.1)
            set_font(p.add_run(text), size=10.5, color=NAVY, bold=cell is row.cells[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    add_callout(document, data["lead"], fill="FFF2C9", color=NAVY)


def add_prompt_card(document, item):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CREAM)
    title = cell.paragraphs[0]
    set_spacing(title, after=4, line=1.1)
    set_font(title.add_run(item["title"]), size=11.5, color=INDIGO, bold=True)
    body = cell.add_paragraph()
    set_spacing(body, after=0, line=1.15)
    set_font(body.add_run(item["text"]), size=10.2, color=NAVY)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build(data, output):
    document = Document()
    configure_document(document)
    numbered = add_numbering(document, "%1.", "decimal")
    numbered_after = add_numbering(document, "%1.", "decimal")
    bullets = add_numbering(document, "•", "bullet", color="4B57E8")

    add_cover(document, data)
    add_heading(document, "Начните с одной задачи", 1)
    for step in data["start_steps"]:
        p = document.add_paragraph()
        apply_numbering(p, numbered)
        set_font(p.add_run(step))
    add_callout(
        document,
        "Не передавайте пароли, банковские данные, фото документов, чужие медицинские данные, переписки и закрытые рабочие материалы.",
        fill="FBE8E6",
        color=CORAL,
    )

    document.add_page_break()
    add_heading(document, "Формула понятного запроса", 1)
    add_body(document, "Хороший запрос не обязан быть идеальным. Достаточно объяснить задачу так, как вы объяснили бы её внимательному помощнику.")
    formula = document.add_table(rows=1, cols=1)
    set_table_geometry(formula, [9360])
    cell = formula.cell(0, 0)
    shade_cell(cell, LAVENDER)
    for index, line in enumerate(data["prompt_formula"]):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_spacing(p, after=4, line=1.15)
        set_font(p.add_run(line), size=10.7, color=NAVY, name="Courier New")
    add_heading(document, "Готовые запросы", 1)
    for item in data["prompts"][:3]:
        add_prompt_card(document, item)

    document.add_page_break()
    add_heading(document, "Ещё три запроса", 1)
    for item in data["prompts"][3:]:
        add_prompt_card(document, item)
    add_callout(document, "После первого ответа продолжайте диалог: «Что здесь может быть неверно?», «Каких данных не хватает?», «Покажи более простой вариант».")

    document.add_page_break()
    add_heading(document, "Проверка перед использованием", 1)
    for item in data["checklist"]:
        p = document.add_paragraph()
        apply_numbering(p, bullets)
        set_font(p.add_run(item))
    add_callout(
        document,
        "Нейросеть готовит черновик. Решение, проверка и ответственность остаются у человека.",
        fill="E7F6F4",
        color=AQUA,
    )
    add_heading(document, "Карта сервисов", 1)
    add_body(document, "Начните с одного универсального помощника. Остальные сервисы подключайте только под конкретную задачу.")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for cell, text in zip(header, ("Сервис", "Когда пригодится")):
        shade_cell(cell, LAVENDER)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1)
        set_font(p.add_run(text), size=10.5, color=NAVY, bold=True)
    for item in data["services"][:3]:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        set_spacing(p, after=2, line=1.1)
        add_hyperlink(p, item["name"], item["url"], bold=True)
        p = cells[1].paragraphs[0]
        set_spacing(p, after=2, line=1.1)
        set_font(p.add_run(item["purpose"]), size=10.1)
        source_p = cells[1].add_paragraph()
        set_spacing(source_p, after=0, line=1)
        add_hyperlink(source_p, "официальная справка", item["source"], color=AQUA)
    set_table_geometry(table, [2300, 7060])

    document.add_page_break()
    add_heading(document, "Карта сервисов — продолжение", 1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for cell, text in zip(header, ("Сервис", "Когда пригодится")):
        shade_cell(cell, LAVENDER)
        p = cell.paragraphs[0]
        set_spacing(p, after=0, line=1)
        set_font(p.add_run(text), size=10.5, color=NAVY, bold=True)
    for item in data["services"][3:]:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        set_spacing(p, after=2, line=1.1)
        add_hyperlink(p, item["name"], item["url"], bold=True)
        p = cells[1].paragraphs[0]
        set_spacing(p, after=2, line=1.1)
        set_font(p.add_run(item["purpose"]), size=10.1)
        source_p = cells[1].add_paragraph()
        set_spacing(source_p, after=0, line=1)
        add_hyperlink(source_p, "официальная справка", item["source"], color=AQUA)
    set_table_geometry(table, [2300, 7060])
    add_body(
        document,
        "Функции, бесплатные лимиты, доступность по странам и условия использования меняются. Перед важной задачей проверьте текущий интерфейс и правила на официальном сайте.",
    )
    add_heading(document, "Что делать после вебинара", 1)
    for text in (
        "Сегодня: выберите одну маленькую задачу и сохраните удачный запрос.",
        "В течение недели: повторите тот же сценарий три раза и отмечайте время, правки и ошибки в журнале.",
        "После третьего повтора: решите, оставить сценарий, изменить его или отказаться.",
    ):
        p = document.add_paragraph()
        apply_numbering(p, numbered_after)
        set_font(p.add_run(text))
    add_callout(
        document,
        "В комплекте также есть рабочая тетрадь и XLSX-журнал трёх повторов. Ссылка на запись и Telegram находится в письме, которым вы получили материалы.",
        fill="FFF2C9",
        color=NAVY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = "Нейросети простым языком — материалы участника"
    document.core_properties.subject = "Памятка, готовые запросы и карта сервисов"
    document.core_properties.author = "Alex Fintore"
    document.core_properties.keywords = "нейросети, вебинар, памятка, запросы, сервисы"
    document.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path("program/materials/participant-starter-guide.json"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("artifacts/participant-kit/neiroseti-prostym-yazykom-starter-guide.docx"),
    )
    args = parser.parse_args()
    data = json.loads(args.source.read_text(encoding="utf-8"))
    build(data, args.output)


if __name__ == "__main__":
    main()
